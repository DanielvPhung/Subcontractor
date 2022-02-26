from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_functions import add_reg_text, add_line_spaces, add_bold_text, bold_run
from datetime import date
import os
from docx2pdf import convert

PATH = os.path.dirname(os.path.realpath(__file__))

def add_title(form):
    title = form.add_paragraph(style = 'Normal')
    run = title.add_run('SUBCONTRACTOR AGREEMENT')
    run.bold = True
    run.font.size = Pt(14) 

    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def add_intro(form , name, address, city_province, postal):
    today = date.today()
    intro_paragraph = form.add_paragraph()
    intro_paragraph.style = 'Normal'

    add_bold_text(intro_paragraph, 'This Subcontract Agreement was created on ' + str(today) + ' by and between the parties: ')
    add_line_spaces(intro_paragraph,2)

    add_bold_text (intro_paragraph, 'First Party: Sinkat Contracting Inc.')
    
    add_reg_text(intro_paragraph,' hereinafter called the "Contractor" and ')
    add_line_spaces(intro_paragraph,2)
    
    add_bold_text(intro_paragraph, ('Second Party: ' + name))

    add_reg_text(intro_paragraph,' hereinafter called the "Subcontractor", who currently resides at: ')
    add_line_spaces(intro_paragraph,1)

    add_reg_text(intro_paragraph,address)
    add_line_spaces(intro_paragraph, 1 )
    
    add_reg_text(intro_paragraph,city_province)
    add_line_spaces(intro_paragraph,1)
    
    add_reg_text(intro_paragraph,postal)
    add_line_spaces(intro_paragraph,2)

    add_reg_text(intro_paragraph,'Both parties, namely ')

    add_bold_text(intro_paragraph,'Sinkat Contracting Inc.')

    add_reg_text(intro_paragraph,' and ' )

    add_bold_text(intro_paragraph,name)
    add_reg_text(intro_paragraph,', hereby confirm their full understanding and agreement to all terms in this agreement as deateiled below: ' )



def add_rules(form, hourly_rate):
    #Change Hourly rate on text file
    
    rules = [
        '“Subcontractor” shall be working as a “General Construction Labourer”, and shall be working in the capacity of an independent contractor and not as an employee of “Contractor”',

        'Subcontractor shall be responsible for completing construction assignments assigned to him, via verbal or written (including email) communications with the contractor on a day-to-day basis.',

        '“Subcontractor” is not required to work for the “Contractor” based on specific work schedule or certain hours.',

        '“Contractor” will pay “Subcontractors based on task completion only and will not pay him fixed salary, neither any taxable benefits, CPP or EI contribution.',

        '“Subcontractor” is totally responsible for completing assigned job in a way that best suites his condition within the limits defined by the “Contractor”.',

        'This agreement shall remain in full force and effective until terminated. Either Party may terminate this Agreement at any time, with or without cause.',

        'Additions or deletions can be made at any time to this agreement by “Contractor” and “Subcontractor” will be provided with a new copy to be review and understood by both parties in a timely manner.',

        'For each construction assignment, “Contractor” agrees to pay “Subcontractor” based on hours spent on job (at a rate of $' + str(hourly_rate) + ' per hour), based on percentage of job completion (measured by square feet or any other relevant measure according to the nature of the job), or based on mixture of both time and percentage of completion. The appropriate method of payment will be determined separately before each job based on the job nature and performance circumstances. ',

        '“Contractor” agrees to pay / reimburse “Subcontractor” additional fees on the following conditions:',

        'a) If “Subcontractor” has to travel more than 40k to a work site using his personal vehicle, based on a rate of 35 cent per kilometer.',

        'b) If “Subcontractor” has to buy work-related supplies or tools up to a daily limit of $500 and provided he provide “Contractor” with the purchase invoices / receipts. Such supplies or tools will be owned by the contractors once reimbursed.',

        'c) Other reimbursement as agreed based on each project individually.',

        '“Contractor” shall be responsible for all cost of supplies needed to preform jobs assigned to Subcontractor.',

        '“Subcontractor” will be responsible for providing all tools and clothing required to perform his assignments, other than unique tools required for some jobs that both parties agreed that “Contractor” should provide.',

        '“Subcontractor” shall be responsible for filing his own taxes with the Canada Revenue Agency, and the “Contractor” based on the fact of his relationship with the subcontractor, which is a “contractor-subcontractor relation” shall not provide the “Subcontractor” with T4 slips, paystubs or record of employment upon contract termination or at calendar year end.',

        '“Contractor” shall not be responsible to reimburse or compensate “Subcontractor” of his personal expense incurred while carrying his duties and executing assignments assigned to him by “Contractor”. Thus, expenses incurred by “Subcontractor” like meals, cell phone expenses, accommodation expenses or driving to work-site expenses shall not be reimbursed, excluding from these expenses mentioned in article “8” above which shall be reimbursed by the “Contractor”',

        '“Subcontractor” shall be fully responsible for the outcome of his misconduct and consequences of his actions and negligence. This shall include liability resulted towards any third party and liability resulted towards clients / other contractors for whom construction job is performed.',

        '“Subcontractor” shall be fully responsible for any breach of laws and regulations, while performing jobs assigned to them.',

        '“Contractor” does not guarantee, nor obligated to ensure that “Subcontractor” shall work for a specific number of hours, or work specific days of the week, or guarantee any minimum amount of work for a single assignment, However, “Subcontractor” has the right to refuse any job or project assigned to him, with no effect on the contractual relation with the “Contractor”.',

        '“Contractor” does not guarantee to “Subcontractor” a specific or average amount of revenue/pay or certain amount or percentage of profit. In fact, “Subcontractor” may even incur losses due to unforeseeable complications with the jobs assigned to them, or due to damages or penalties caused by their poor performance of jobs assigned to them, or as a result of negligence or failure to adhere to regulatory.',

        'Any revenue “Subcontractor” generates or receive from a client or other contractor whom the “Contractor” is performing a job and or has an agreement with, while executing his assignment or job to the same client or contractor should be paid back to the “Contractor”.',

        '“Subcontractor” hereby agrees and confirm that he shall not compete in any manner with “Contractor”, though approaching or generating business or earn revenue from clients/contractors with whom the “Contractor” has a business relation while this contract in valid and 3 years after the termination of this contract.',

        'After each job or project assigned to “Subcontractor”, “Subcontractor” is responsible to email/fax all the invoices to “Contractor” within 3 business days after job end.',

        '“Subcontractor” agrees to follow and uphold all regulations set by the Construction industries norms and standards while performing jobs assigned to him.',

        '“Subcontractor” understand that “Contractor” has a zero-tolerance policy and does not permit anything that is deemed contraband or illegal by the Federal or provincial laws of Canada and accept full responsibility and penalties due to his own actions regarding this policy.',

        '“Subcontractor” shall not take work related equipment and tools owned by “Contractor” anywhere other than agreed upon locations and destinations and shall not use those equipment and tools for personal use without “Contractor” approval.',

        'The work performed by the “Subcontract” shall be at the risk of the “Subcontractor” exclusively. Subcontractor hereby indemnifies and holds “Contractor” corporation, its shareholders and directors as harmless from and against any and all claims, actions, losses, judgments, or expenses, including reasonable attorney’s fees, arising from or in any way connected with the work performed, or services provided to “Contractor” during the term of this Agreement. ',

        '“Subcontractor” should submit to “Contractor” within 24 hours after finishing each job / project a log sheet that shows hours and materials dedicated to the job, with relevant invoices or receipts of purchases made by “Subcontractor”, along with photos, full description of performance, and detailed notes of any issue related to performance of job assigned.',
    ]

    rule_count = 1
    for rule in rules:
        #Strip whitespaces from front and back 
        rule = rule.rstrip()
        rule = rule.lstrip()
        
        #Ignore Empty Lines
        if(rule == ''):
            pass

        #If we have a sublist, for example a) b) and c)
        elif(rule[1] == ')'):
            sublist = form.add_paragraph(style = 'List')
            sublist.paragraph_format.left_indent = Inches(0.75)
            run = sublist.add_run(rule)
            run.add_break()
            
        #Insert Line, increment our rule counter
        else: 
            paragraph = form.add_paragraph(style = 'List Number')
            paragraph.paragraph_format.left_indent = Inches(0.5)
            run = paragraph.add_run(rule)
            run.add_break()
            rule_count+=1

def add_signatures(form, email, phone_number, sin):
    line_seperator = form.add_paragraph('__________________________________________________')    
    line_seperator_format = line_seperator.paragraph_format
    line_seperator_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    paragraph = form.add_paragraph(style = 'Normal')
    add_reg_text(paragraph, ('This Agreement, including all terms and conditions hereof, is expressly agreed to and constitutes the ' +
                             'entire Agreement as of this date. No other Agreement or understandings, verbal or written, expressed ' +
                             'or implied, are a part of this Agreement unless specified herein. '))

    add_line_spaces(paragraph, 2)
    add_bold_text(paragraph, '[Contractor Signature]')
    add_line_spaces(paragraph, 3)
    add_bold_text(paragraph, '[Subcontractor Signature]')
    add_line_spaces(paragraph, 3)
            

    table = form.add_table(rows = 4, cols = 2)
    for i in range(0, len(table.columns)):
        curr_cell = table.cell(0,i) 
        curr_cell.text = 'Company: __________________________'
        bold_run(curr_cell.paragraphs[0].runs[0] )
        
        curr_cell = table.cell(1,i) 
        curr_cell.text = 'Name/Title: ________________________'
        bold_run(curr_cell.paragraphs[0].runs[0] )

        curr_cell = table.cell(2,i) 
        curr_cell.text = 'Signature: _________________________'
        bold_run(curr_cell.paragraphs[0].runs[0] )

        curr_cell = table.cell(3,i) 
        curr_cell.text = 'Date: _____________________________'
        bold_run(curr_cell.paragraphs[0].runs[0] )


    #Subcontractor Info Paragraph
    sc_info = form.add_paragraph(style = 'Normal')
    add_line_spaces(sc_info, 2)

    add_bold_text(sc_info, ('Email Address: ' + email))
    add_line_spaces(sc_info, 2)

    add_bold_text(sc_info, ('Phone Number: ' + phone_number))
    add_line_spaces(sc_info, 2)

    add_bold_text(sc_info, ('SIN: ' + sin))
    add_line_spaces(sc_info, 2)



def create_sc_form(name, sin, address, city_province, postal_code, email, phone_number):

    DOCUMENT_NAME_DOCX = name + ' Subcontractor Form.docx'
    DOCUMENT_NAME_PDF = name + ' Subcontractor Form.pdf'

    form = Document()
    style = form.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(11)

    hourly_rate = input('Enter Hourly Rate: ')
    add_title(form)
    add_intro(form ,name, address, city_province, postal_code)
    add_rules(form , hourly_rate)
    add_signatures(form, email, phone_number,sin)


    form.save(PATH + '\\' + DOCUMENT_NAME_DOCX)
    convert(PATH + '\\' + DOCUMENT_NAME_DOCX)
    print('PDF successfully created')

#Testing Purposes
def main():
    create_sc_form('Name', '999999999', 'address' , 'city_province', 'postal_code', 'email', '(999) 999 9999')
    input('Complete')
if __name__== "__main__" :
    main()